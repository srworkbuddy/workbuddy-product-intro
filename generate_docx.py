#!/usr/bin/env python3
"""生成 WorkBuddy 产品介绍 DOCX"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# Colors
BLUE = RGBColor(0x00, 0x52, 0xD9)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x66, 0x66, 0x66)
GREEN = RGBColor(0x00, 0x88, 0x44)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLUE_BG = "E8F0FE"
GREEN_BG = "F0FFF5"

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = BLUE
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def add_heading2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = DARK
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def add_heading3(doc, text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = BLUE
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def add_body(doc, text, bold=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.bold = bold
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(4)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(3)
    return p

def add_numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(3)
    return p

def add_highlight(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = BLUE
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.italic = True
    pf = p.paragraph_format
    pf.left_indent = Cm(1)
    pf.space_before = Pt(8)
    pf.space_after = Pt(8)
    # Add left border
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="12" w:space="8" w:color="0052D9"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p

def add_green_box(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0x44)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    pf = p.paragraph_format
    pf.left_indent = Cm(0.5)
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    # Green background via shading
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{GREEN_BG}" w:val="clear"/>')
    pPr.append(shd)
    return p

def add_material_links(doc, links):
    p = doc.add_paragraph()
    run = p.add_run("🔗 素材链接")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = BLUE
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    for name, url in links:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(1)
        p2.paragraph_format.space_after = Pt(2)
        p2.paragraph_format.space_before = Pt(2)
        # Bullet character
        run0 = p2.add_run("• ")
        run0.font.size = Pt(9)
        run0.font.color.rgb = GRAY
        # Hyperlink
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        part = doc.part
        r_id = part.relate_to(url, "hyperlink", is_external=True)
        hyperlink = parse_xml(
            f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" {nsdecls("r")}/>'
        )
        new_run = parse_xml(
            f'<w:r {nsdecls("w")}>'
            f'<w:rPr><w:rStyle w:val="Hyperlink"/><w:sz w:val="18"/>'
            f'<w:color w:val="0052D9"/><w:u w:val="single"/>'
            f'<w:rFonts w:ascii="Microsoft YaHei" w:eastAsia="Microsoft YaHei"/></w:rPr>'
            f'<w:t xml:space="preserve">{name}</w:t></w:r>'
        )
        hyperlink.append(new_run)
        p2._p.append(hyperlink)

def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        set_cell_shading(cell, "0052D9")
    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9.5)
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            if c_idx == 0:
                run.font.bold = True
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()  # spacer
    return table

def add_enterprise_scene(doc, title, pain_point, steps, customer_value, links):
    add_heading3(doc, title)
    add_body(doc, f"痛点：{pain_point}")
    add_body(doc, "实践路径：", bold=True)
    for step in steps:
        add_numbered(doc, step)
    add_green_box(doc, f"客户价值：{customer_value}")
    add_material_links(doc, links)
    doc.add_paragraph()  # spacer


def main():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # ===== Cover =====
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("WorkBuddy")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("全场景职场 AI 原生智能体桌面工作台")
    run.font.size = Pt(16)
    run.font.color.rgb = GRAY

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("产品介绍 · 场景实践 · 使用案例")
    run.font.size = Pt(12)
    run.font.color.rgb = DARK

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("一句话让 AI 替你上班")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = GREEN

    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("腾讯 CSIG · 2026年5月")
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY

    doc.add_page_break()

    # ===== 1. 产品概述 =====
    add_heading1(doc, "一、产品概述")
    add_heading2(doc, "1.1 产品定位")
    add_body(doc, 'WorkBuddy 是腾讯推出的桌面级 AI 助手，被媒体叫作\u201c腾讯版小龙虾\u201d。跟只能在对话框里回答问题的传统 AI 不一样，WorkBuddy 能直接操作你电脑上的文件——你说一句话，它自己去规划、执行，最后把做好的东西交给你。2026 年 3 月 9 日上线。')
    add_highlight(doc, "产品于 2026 年 3 月 9 日正式上线，基于腾讯 CodeBuddy 同源架构打造")

    add_heading2(doc, "1.2 核心差异化")
    add_body(doc, "传统 AI 聊天工具 vs WorkBuddy 的本质区别：", bold=True)
    make_table(doc,
        ["对比维度", "传统 AI 聊天工具", "WorkBuddy 智能体工作台"],
        [
            ["它的角色", "写作顾问 / 建议提供者", "办公执行者 / 任务完成同事"],
            ["交付结果", "方法建议、模板、润色建议", "分类结果、回复草稿、正式邮件、可发送交付物"],
            ["能否接触现场", "不能，需用户复制粘贴", "能读取文件/邮箱/表格，理解上下文"],
            ["用户要做什么", "自己筛选、判断、写作、发送", "提出目标，检查结果，确认发送"],
            ["核心差别", "让员工\u201c知道怎么处理\u201d", "让员工\u201c直接拿到处理后的结果\u201d"],
        ],
        [3, 5, 5]
    )

    # ===== 2. 核心能力 =====
    doc.add_page_break()
    add_heading1(doc, "二、核心能力")

    add_heading2(doc, "2.1 智能任务自动化执行")
    add_body(doc, "你说一句话，它自己把任务拆开、排好步骤、跑完全程。适合周期性、重复性的办公活儿。")
    add_bullet(doc, "自动生成日报/周报/季度报告")
    add_bullet(doc, "定时抓取竞品信息和行业动态")
    add_bullet(doc, "整理会议录音与纪要")

    add_heading2(doc, "2.2 本地文件全流程操作")
    add_body(doc, "你授权之后，它可以直接读写你电脑上的文件——批量编辑、分类整理、格式转换、归档。")
    add_bullet(doc, "整理桌面文件按类型分类")
    add_bullet(doc, "批量重命名、分类归档")
    add_bullet(doc, "处理 Excel 表格、编辑 Word 文档")

    add_heading2(doc, "2.3 多模态办公内容创作")
    add_body(doc, "写文案、做 PPT、画海报、出数据图表，风格、框架、细节随你调。")

    add_heading2(doc, "2.4 专业数据处理与报表生成")
    add_body(doc, "数据清洗、统计分析、图表可视化、自动出报表。表格数据跑一遍，核心指标和结论就出来了。")

    add_heading2(doc, "2.5 跨平台远程电脑控制")
    add_body(doc, "支持 MCP 协议，连接企业微信、QQ、飞书、钉钉等办公平台。手机上发一句话，办公室电脑就开始干活，做完把结果给你。")
    add_highlight(doc, "企微/QQ 三步连上，手机发指令，电脑端自动干活")

    add_heading2(doc, "2.6 Skills 技能包系统")
    add_body(doc, "自带 20 多个 Skills 技能包，不用写代码就能扩展新能力。日常操作可以封装成可复用的技能，也能在团队里分享、上传到开源社区。")
    add_bullet(doc, "办公效率全家桶（Word、Excel、PPT、PDF 处理）")
    add_bullet(doc, "内容创作与多平台适配（小红书、公众号、SEO）")
    add_bullet(doc, "垂直行业专业任务（法律、金融、人力资源）")

    add_heading2(doc, "2.7 多模型自由切换")
    add_body(doc, "不绑定单一模型，五大主流大模型随意切：")
    make_table(doc,
        ["模型", "特点", "适用场景"],
        [
            ["腾讯混元", "强大中文理解、多模态处理", "报告撰写、文案生成、数据分析"],
            ["DeepSeek-V3/R1", "推理能力强、精度高", "复杂架构设计、调试任务"],
            ["GLM-5 系列", "中文语义理解优秀", "中文语境复杂语义处理"],
            ["Kimi-K2 系列", "超长上下文（256K）", "超长文档阅读、多文档分析"],
            ["MiniMax-M2 系列", "多媒体内容生成", "内容创作、角色扮演"],
        ],
        [3, 5, 5]
    )

    # ===== 3. 企业版飞轮 =====
    doc.add_page_break()
    add_heading1(doc, "三、企业版核心逻辑：AI 工作飞轮")
    add_highlight(doc, "企业版不是给员工多一个 AI 聊天框。它跑的是一条飞轮：统一入口 → 超级个体 → 超级团队 → 企业 AI 资产 → 企业数字员工 → 组织生产力持续释放。")
    add_material_links(doc, [("WorkBuddy企业版主线逻辑说明材料", "https://csig.lexiangla.com/pages/dd473bab87ae46ab91a1049a101d35b4")])

    add_heading2(doc, "3.1 五大核心概念")
    make_table(doc,
        ["概念", "定义", "价值"],
        [
            ["超级个体", "员工用 WorkBuddy 调用组织能力干活", "个人能力直接放大，普通员工也能做更专业的活儿"],
            ["超级团队", "人 + 数字员工组队干活", "团队围绕任务持续推进，不只是沟通"],
            ["企业 AI 资产", "每次任务沉淀知识、技能、经验和案例", "经验从个人释放，变成组织可复用能力"],
            ["企业数字员工", "7×24 在线，懂岗位、有技能、熟悉企业知识", "数字员工越来越懂企业，能承担更多专业任务"],
            ["统一治理", "管理员统一管控智能体的权限、安全和审计", "企业敢把更多流程交给 AI，形成规模化落地"],
        ],
        [3, 5, 5]
    )

    add_heading2(doc, "3.2 Agent Suite 企业套件")
    make_table(doc,
        ["套件", "定位", "价值"],
        [
            ["腾讯文档", "内容生产与协作底座", "AI 生成的东西直接进协作体系"],
            ["腾讯网盘", "文件和素材管理底座", "本地产物、团队共享文件都能存、能回溯"],
            ["腾讯乐享", "知识治理底座", "制度、培训、经验变成数字员工的知识"],
        ],
        [3, 5, 5]
    )

    # ===== 4. 使用场景 =====
    doc.add_page_break()
    add_heading1(doc, "四、使用场景")

    # 4.1 办公场景
    add_heading2(doc, "4.1 办公场景：高效信息沟通")
    add_heading3(doc, "场景描述")
    add_body(doc, "周一早上 9 点，员工打开邮箱发现 38 封未读邮件，包含领导催进度、客户报价、HR 表格、项目协作、会议通知等不同类型事项。")
    add_heading3(doc, "WorkBuddy 做法步骤")
    add_numbered(doc, "读取未读邮件汇总文件，识别每封邮件的业务含义和处理责任")
    add_numbered(doc, '将 38 封邮件分为\u201c必须亲自回/助理代回/仅知会/可忽略\u201d4 类')
    add_numbered(doc, '为\u201c必须亲自回\u201d的邮件生成专业中文回复草稿')
    add_numbered(doc, "处理重点英文客户邮件：先生成中文审稿，再翻译成商务英语版")
    add_numbered(doc, "员工确认后可继续完成发送或回写")
    add_heading3(doc, "提效效果")
    add_body(doc, "传统方式 2-3 小时 → WorkBuddy 10-20 分钟，员工从执行者变成审核者。")
    add_material_links(doc, [
        ("WorkBuddy办公场景01_高效信息沟通_最佳实践案例", "https://csig.lexiangla.com/pages/e769ffa538b64e888380bf1659f58e9e"),
        ("传统AI聊天_vs_WorkBuddy干活_最佳实践案例", "https://csig.lexiangla.com/pages/b96442a34bef4915b0e42ac078116d29"),
    ])

    # 4.2 PM 场景
    add_heading2(doc, "4.2 PM 场景：项目过程管理")
    make_table(doc,
        ["序号", "场景", "典型输入", "标准输出", "直接价值"],
        [
            ["1", "会议纪要闭环", "会议记录/聊天记录", "结论/行动项/责任人/DDL/风险项", "会后闭环速度提升"],
            ["2", "周报/月报生成", "本周事项、风险、里程碑", "管理层口径周报草稿", "写作时间下降、口径更统一"],
            ["3", "验收材料打包", "验收报告、清单、截图", "可提交材料包+检查清单", "减少漏项、提升交付质量"],
            ["4", "项目知识检索", "历史PRD、纪要、周报、合同", "关键结论与出处", '减少\u201c找资料\u201d耗时'],
            ["5", "远程触发任务", "微信/企微/QQ 指令", "周报草稿、资料归档结果", "碎片时间也能推进项目"],
        ],
        [1.5, 3, 4, 4, 3.5]
    )
    add_material_links(doc, [
        ("WorkBuddy-PM在日常工作应用场景实践-v5", "https://csig.lexiangla.com/pages/41aa17715ff64c15a8bd662ab8963cff"),
        ("TAPD + 乐享 + WorkBuddy 联动方案", "https://csig.lexiangla.com/pages/98928c2886ac42e58d289bd6c8878a8d"),
    ])

    # 4.3 全岗位高频场景
    doc.add_page_break()
    add_heading2(doc, "4.3 全岗位高频场景")
    add_body(doc, "来自 WorkBuddy 官方公众号和腾讯云开发者社区的真实用户场景，覆盖行政、HR、财务、运营、法务、教育等岗位。")

    # HR
    add_enterprise_scene(doc,
        "HR · 简历筛选与招聘提效",
        "逐份打开简历手动比对JD、填表打分干到眼花；招聘文案从头写，入职通知一份份发",
        ["批量读取候选人简历，按匹配度自动打分，排出面试优先级",
         "生成《候选人匹配度评分报告》，含排名、核心亮点、关键提醒",
         "自动生成岗位招聘文案（岗位职责/任职要求/薪资范围）",
         "5分钟批量发送100份个性化入职通知"],
        "简历筛选半天→15分钟 | 招聘文案一键生成 | 入职通知批量个性化发送",
        [("8大高频办公场景提示词", "https://ai-bot.cn/workbuddy-office-prompts/"),
         ("WorkBuddy最全干货：6大岗位场景", "https://www.163.com/dy/article/KP3O2S2905383FJC_pdya11y.html")])

    # 财务
    add_enterprise_scene(doc,
        "财务 · 发票报销与报表汇总",
        "月底一张张拍照录入发票，金额、日期、项目名错一个字整张表重来；多张Excel报表汇总靠手动",
        ["批量导入发票文件，自动提取金额/日期/项目名称/代码",
         "按\u201c客户名字+开具时间+总金额\u201d命名归档",
         "45张发票5分钟生成报销明细表",
         "多个Excel报表自动汇总，按月/季度生成可视化图表"],
        "发票录入2小时→5分钟 | 报销明细可直接提交 | 财务报表自动汇总",
        [("8大高频办公场景提示词", "https://ai-bot.cn/workbuddy-office-prompts/"),
         ("腾讯WorkBuddy上线！AI办公搭子来袭", "https://cloud.tencent.com/developer/article/2638060")])

    # 行政
    add_enterprise_scene(doc,
        "行政 · 会议纪要与文件归档",
        "会议录音转写后还要手动提炼决议和行动项；文件命名混乱、分类归档靠手工；各类通知格式不统一",
        ["会议录音/零散笔记进入WorkBuddy，自动提炼关键决策、行动项、责任人和截止时间",
         "生成标准化会议纪要，按部门分类",
         "办公文件自动归档，按类别（合同/通知/报表）分类，统一重命名为\u201c日期+文件名\u201d",
         "批量生成各类通知文件"],
        "纪要整理30分钟→3分钟 | 文件归档自动化 | 通知格式统一规范",
        [("WorkBuddy最全干货：6大岗位场景", "https://www.163.com/dy/article/KP3O2S2905383FJC_pdya11y.html"),
         ("腾讯WorkBuddy上线！AI办公搭子来袭", "https://cloud.tencent.com/developer/article/2638060")])

    # 运营/市场
    add_enterprise_scene(doc,
        "运营/市场 · 多平台内容与竞品分析",
        "搜资料1小时、组织语言1小时、排版半小时，交稿还被说\u201c不够网感\u201d；竞品分析靠手动搜集",
        ["输入主题自动搜集最新资料，15分钟出推文初稿",
         "多平台文案一键生成（公众号/小红书/短视频脚本），支持风格定制",
         "分析竞品推广策略，整理成复盘报告，指出自身优势与改进方向",
         "公众号周报自动化：每周一自动拉取阅读数据分析Top3/Bottom3"],
        "15分钟完成3篇小红书图文 | 公众号效率提升300%+ | 竞品分析半天→30分钟",
        [("我用WorkBuddy打造公众号内容工厂", "https://cloud.tencent.com/developer/article/2640886"),
         ("8大高频办公场景提示词", "https://ai-bot.cn/workbuddy-office-prompts/")])

    # 法务
    add_enterprise_scene(doc,
        "法务 · 合同审查与风险排查",
        "老板给了10份合同逐一审查风险，每次只能上传1份；合同条款多，关键风险点容易遗漏",
        ["批量上传合同文件，给一条统一指令，系统自动逐个处理",
         "自动识别合同中的风险条款、异常约定和合规问题",
         "生成风险审查报告，标注关键提醒和修改建议",
         "合同版本对比，自动标注差异条款"],
        "10份合同审查1天→2小时 | 风险点自动标注 | 版本差异一目了然",
        [("别只当普通工具用！WorkBuddy 5个隐藏功能", "https://m.toutiao.com/group/7620648693502312996/")])

    # 教育
    add_enterprise_scene(doc,
        "教育 · 出题/成绩分析/课件制作",
        "出阅读理解题太费时间，网上题目千篇一律；考完试成绩分析要用Excel算半天；做一个课件至少1小时",
        ["根据课文自动生成阅读理解题（选择+简答+参考答案），导出Word直接打印",
         "成绩表自动分析：按班级统计各科平均分/最高分/及格率/优秀率，生成柱状图",
         "Word文档一键生成课件PPT",
         "复旦大学200+师生现场体验\u201cAI百校行\u201d"],
        "出题1小时→5分钟 | 成绩分析30分钟→30秒 | 课件1小时→20分钟",
        [("用了一周WorkBuddy，老师必备的效率神器", "https://www.yeyulingfeng.com/535759.html"),
         ("高中数学老师的AI实战记录", "https://cloud.tencent.com.cn/developer/article/2663569"),
         ("高校\u201c龙虾\u201d专列开进复旦", "https://news.qq.com/rain/a/20260410A04DN700")])

    # 4.4 企业版场景实践
    doc.add_page_break()
    add_heading2(doc, "4.4 企业版场景实践")
    add_body(doc, "来自 WorkBuddy 企业版主线逻辑，非乐享联动的独立企业场景。")

    add_enterprise_scene(doc,
        "场景一：会议到任务闭环",
        "会议纪要有了，但行动项没人跟，责任人和后续进展容易断，会后闭环全靠人追",
        ["会议记录、聊天记录、语音转写进入 WorkBuddy",
         "数字员工提炼结论、行动项、责任人、DDL 和风险项",
         "Teams 把行动项分配给成员和数字员工",
         "自动化任务定期提醒、跟踪和汇总进展",
         "后续产出、复盘和决策沉淀为团队资产",
         "下一次会议可直接复用这些组织记忆"],
        "会议不再止于纪要 | 行动项持续推进 | 任务状态可见 | 复盘资料可沉淀 | 团队协同从沟通走向执行",
        [("WorkBuddy-PM在日常工作应用场景实践-v5", "https://csig.lexiangla.com/pages/41aa17715ff64c15a8bd662ab8963cff")])
    add_body(doc, '实操建议：固定模板先行，字段保持一致便于追踪。会后 10 分钟内完成\u201c纪要+分派+确认\u201d')

    add_enterprise_scene(doc,
        "场景二：管理层周报 / 例会 / 经营复盘",
        "信息散落在会议纪要、项目材料、沟通记录里，每周花大量时间收集整理，周报和复盘高度依赖手工搬运",
        ['乐享沉淀日常经营类内容：项目纪要、复盘文档、制度更新、风险记录、业务分析、运营日报/周报素材',
         '管理者通过 WorkBuddy 发起：\u201c整理本周重点事项和风险点\u201d、\u201c生成销售周报\u201d、\u201c汇总项目例会待办\u201d',
         'WorkBuddy 生成：周报、例会提纲、跟进行动清单、经营摘要、风险项列表',
         '接入企微后支持结果推送和远程办公场景'],
        "节省管理层和助理整理时间 | 强化经营信息统一口径 | 提高例会和复盘效率",
        [("WorkBuddy + 乐享AI知识库联合方案", "https://csig.lexiangla.com/pages/0be95da3d040443bacab438f7c678c52")])

    # 4.5 乐享知识库联动场景
    doc.add_page_break()
    add_heading2(doc, "4.5 乐享知识库联动场景")
    add_body(doc, "WorkBuddy + 腾讯乐享知识库联合方案，三层架构（底层乐享知识底座 \u2192 中层AI理解与调用 \u2192 上层WorkBuddy任务执行与结果交付）。乐享负责知识沉淀与权限治理，WorkBuddy 负责智能调用与任务执行。")

    add_enterprise_scene(doc,
        "场景一：售后 / 客服问题处理",
        "客服培训慢，复杂问题升级多，历史案例和知识库难以复用，新人不会查也不知道哪个答案是对的",
        ["售后同学每天打开 WorkBuddy 处理客户问题",
         "调用乐享知识库中的 FAQ、故障案例、操作手册和话术模板",
         "售后数字员工整理可能原因、排查顺序和客户回复草稿",
         "需要升级时，数字员工加入 Teams，与产品、研发和二线支持协同",
         "最终回复、排查记录和新案例沉淀回企业知识库",
         "沉淀的案例和 SOP 再赋能下一次售后数字员工"],
        "客户回复更快 | 排查路径更标准 | 新人更容易上手 | 案例持续沉淀 | 售后经验从个人能力变成组织能力",
        [("WorkBuddy企业版主线逻辑说明材料", "https://csig.lexiangla.com/pages/dd473bab87ae46ab91a1049a101d35b4"),
         ("WorkBuddy + 乐享AI知识库联合方案", "https://csig.lexiangla.com/pages/0be95da3d040443bacab438f7c678c52")])

    add_enterprise_scene(doc,
        "场景二：销售 / 售前材料快速响应",
        "客户需求一来资料找不到或版本不统一，不同销售讲法不一致，售前专家被反复拉来回答基础问题",
        ["乐享统一沉淀：产品资料、行业方案、成功案例、招投标模板、常见异议回复、实施边界与报价规则",
         "销售在企微或桌面侧向 WorkBuddy 发起任务",
         "WorkBuddy 调用乐享知识库内容后，输出：客户方案初稿、沟通邮件、拜访前准备提纲、行业案例摘要、会后纪要与跟进清单",
         "团队在 Teams 中协同修改，产品、售前和销售共同审阅",
         "定稿保存到腾讯文档，证据来源沉淀到销售弹药库",
         '乐享负责\u201c找准内容\u201d，WorkBuddy 负责\u201c变成输出\u201d'],
        "客户响应速度提升 | 销售口径一致性提高 | 方案制作时间缩短 | 销售输出建立在企业知识基础上",
        [("WorkBuddy + 乐享AI知识库联合方案", "https://csig.lexiangla.com/pages/0be95da3d040443bacab438f7c678c52")])

    add_enterprise_scene(doc,
        "场景三：企业知识运营",
        "企业不缺知识，但知识散在文档、网盘、乐享、业务系统和群聊里，员工找不到、找不准，对合规敏感岗位担心 AI 乱答",
        ["员工在 WorkBuddy 中直接询问制度、流程、产品资料和项目经验",
         "WorkBuddy 通过 Agent Suite 调用腾讯文档、腾讯网盘和腾讯乐享",
         "乐享负责知识治理，确保知识可信、可控、可追溯（权限、有效期、版本管理）",
         "数字员工根据企业知识完成问答、总结、分析和材料生成",
         "新问题、新回答、新材料和知识缺口沉淀为企业 AI 资产"],
        "知识找得到 | 回答答得准 | 权限管得住 | 结果可追溯 | 企业知识越用越好",
        [("WorkBuddy + 乐享AI知识库联合方案", "https://csig.lexiangla.com/pages/0be95da3d040443bacab438f7c678c52")])

    add_enterprise_scene(doc,
        "场景四：新人培训与岗位辅导",
        "培训资料很多但不好找不成体系，培训结束后员工仍不会实操，带教大量依赖老员工一对一答疑",
        ["乐享沉淀岗位 SOP、流程规范、制度说明、典型案例、常见错误、模板和标准口径、培训资料与题库",
         "新人遇到问题，通过知识库快速找到答案",
         "WorkBuddy 进一步生成：学习清单、每周学习计划、错题/易错点总结、岗位周报、导师跟进建议",
         "把培训从静态课件变成动态辅导",
         "把组织经验留在系统里，而不是留在个人脑子里"],
        "缩短新人上岗周期 | 降低培训和带教成本 | 培训从静态课件变动态辅导",
        [("WorkBuddy + 乐享AI知识库联合方案", "https://csig.lexiangla.com/pages/0be95da3d040443bacab438f7c678c52")])

    add_enterprise_scene(doc,
        "场景五：移动办公 / 出差办公助手",
        "人在外面资料在公司电脑里，时间碎片化需要用手机快速处理工作",
        ["企业知识统一沉淀在乐享中，供检索和调用",
         "人在外部通过企微向 WorkBuddy 下发任务",
         "WorkBuddy 调用知识并在本地电脑完成操作，再把结果传回企微",
         '示例：手机企微发起\u201c把客户X的最新案例、方案模板和会后纪要发给我，生成跟进邮件\u201d \u2192 系统调用乐享资料 \u2192 WorkBuddy 整理邮件并处理文件 \u2192 结果直接推回企微'],
        '移动办公效率提升 | 减少跨设备跨人协调成本 | \u201c找资料-处理-回传\u201d一次做完',
        [("WorkBuddy + 乐享AI知识库联合方案", "https://csig.lexiangla.com/pages/0be95da3d040443bacab438f7c678c52")])

    add_enterprise_scene(doc,
        "场景六：制度 / 合规 / 专业支持",
        "内容是否准确、权限是否可控、版本是否最新、使用过程是否可追溯",
        ["乐享沉淀：制度规范、合规要求、审批规则、专业模板、历史案例、审核意见、业务指引",
         "通过权限和有效期控制，保证不同角色看到的是正确、有效、可授权的内容",
         "WorkBuddy 在受控知识基础上生成：制度答复草稿、审核清单、材料整理初稿、内部通知、执行提醒",
         "AI 使用建立在合规治理基础上"],
        "减少错用制度和版本错误 | 提高专业岗位工作规范性 | AI 使用建立在合规治理基础上",
        [("WorkBuddy + 乐享AI知识库联合方案", "https://csig.lexiangla.com/pages/0be95da3d040443bacab438f7c678c52")])

    # ===== 5. 使用案例 =====
    doc.add_page_break()
    add_heading1(doc, "五、使用案例")

    # 案例01
    add_heading2(doc, "案例01：构建个人学术信息助手")
    add_body(doc, "搜广推领域算法人员，用 WorkBuddy 从 0 到 1 搭了一套 arXiv 论文每日自动筛选、精排、深度解读、报告生成、知识库沉淀的全流程自动化系统。三路并行召回（关键词+Embedding 语义+高频作者追踪），召回率 97.4%。每个工作日早上 9 点自动生成 10 篇 9 维结构化精读 + 20 篇速览 + 趋势分析，自动上传乐享知识库并发邮件通知。搭建 2 天，代码约 5000 行全部由 WorkBuddy 生成，外部 API 成本为零。")
    add_green_box(doc, "搭建~2天 | 代码~5000行 | 召回率97.4% | 外部API成本¥0 | 每日阅读5分钟")
    add_material_links(doc, [("WorkBuddy 提效案例：构建个人学术信息助手", "https://csig.lexiangla.com/pages/e2cf5cbc5a8f44d28176829c92dbbe6e")])

    # 案例02
    add_heading2(doc, "案例02：AI 简历筛选提效 90%")
    add_body(doc, "新加坡团队用 WorkBuddy 做 AI 简历筛选：输入岗位需求和候选人简历文件，WorkBuddy 自动解析简历内容、按岗位要求逐项匹配评分、生成筛选结果和推荐理由。人工筛选时间从几小时压到几分钟，筛选的一致性和准确度也提高了，避免了不同 HR 主观判断导致的标准偏差。")
    add_green_box(doc, "提效90% | 人工筛选几小时→几分钟")
    add_material_links(doc, [("\u3010案例\u3011新加坡-AI简历筛选提效90%", "https://csig.lexiangla.com/pages/bfce4bba7e754e38826b50396ced359d")])

    # 案例03
    add_heading2(doc, "案例03：五粮液全民 AI 办公提效动员方案")
    add_body(doc, "基于乐享、企微、WorkBuddy 等工具，给五粮液领导层做了一套全民 AI 办公提效动员方案，从知识管理到智能办公整个链路跑通了。方案覆盖白酒行业典型办公场景：知识检索、会议纪要闭环、周报自动生成、质检提效等。WorkBuddy 在白酒行业的第一个客户案例，验证了制造业+AI办公的落地路径。")
    add_green_box(doc, "白酒行业首个客户落地 | 知识管理→智能办公全链路跑通")
    add_material_links(doc, [
        ("基于乐享、企微、workbuddy等生成五粮液全民AI办公提效动员方案", "https://csig.lexiangla.com/pages/2b1e103954414fedad7c66094930b402"),
        ("五粮液：AI助力质检提效解决方案", "https://csig.lexiangla.com/pages/13fe256cb6184810a790fc1e045247b1"),
    ])

    # 案例04
    add_heading2(doc, "案例04：冠盛客户 WorkBuddy 场景沟通实践")
    add_body(doc, "冠盛客户团队和 WorkBuddy 团队做了一次场景沟通，把制造业日常办公的痛点和需求梳理了一遍。围绕售后问题处理、销售材料快速响应、新人培训、管理层周报复盘等场景逐一验证，确认了 WorkBuddy 在制造业场景能落地。沟通结论：WorkBuddy + 乐享联合方案可以覆盖制造企业核心办公痛点。")
    add_green_box(doc, "制造业场景验证 | 售后/销售/培训/管理场景逐一确认")

    # 案例05
    add_heading2(doc, "案例05：WorkBuddy + 乐享 AI 知识库联合方案落地")
    add_body(doc, 'WorkBuddy 与乐享 AI 知识库联合方案已在多个客户场景验证：乐享负责知识沉淀、权限治理、搜索问答和知识运营，WorkBuddy 负责桌面 AI 智能体、本地执行、企微接入和结果回传。两者组合后形成\u201c知识获取+任务执行+结果交付\u201d完整链路。已覆盖六大场景：销售/售前快速响应、客服/客成/实施支持、新人培训与岗位辅导、管理层周报/例会/经营复盘、制度/合规/专业支持、移动办公/出差办公助手。')
    add_green_box(doc, "6大场景方案 | 知识获取+任务执行+结果交付完整链路")
    add_material_links(doc, [("WorkBuddy + 乐享AI知识库联合方案", "https://csig.lexiangla.com/pages/0be95da3d040443bacab438f7c678c52")])

    # 案例06
    add_heading2(doc, "案例06：PM 项目过程管理全场景提效")
    add_body(doc, '针对项目 PM 的 5 个高频场景（会议纪要闭环、周报/月报生成、验收材料打包、项目知识检索、移动端远程触发），WorkBuddy 提供了从输入到输出的标准化实操路径。核心思路：PM 从\u201c从零写报告\u201d转为\u201c审核和定稿\u201d，从\u201c找资料靠记忆\u201d变成\u201c按出处检索\u201d，从\u201c等回工位再做\u201d变成\u201c先启动任务回位后收口\u201d。建议先从纪要闭环+周报生成开始试点。')
    add_green_box(doc, "5大PM高频场景 | PM角色从执行者→审核者")
    add_material_links(doc, [
        ("WorkBuddy-PM在日常工作应用场景实践-v5", "https://csig.lexiangla.com/pages/41aa17715ff64c15a8bd662ab8963cff"),
        ("TAPD + 乐享 + WorkBuddy 联动方案", "https://csig.lexiangla.com/pages/98928c2886ac42e58d289bd6c8878a8d"),
    ])

    # ===== 6. 数据安全 =====
    doc.add_page_break()
    add_heading1(doc, "六、数据安全与合规")

    add_heading2(doc, "6.1 三重安全机制")
    add_bullet(doc, "本地执行：所有文件处理在本地完成，原始文件从不离开用户电脑")
    add_bullet(doc, "权限隔离：沙箱与容器隔离运行，网关防护，暴露面收敛，传输加密，高危操作拦下来")
    add_bullet(doc, "审计留痕：感知可信过滤、提示词攻击防护、推理过程可视化，每个动作都留痕、可控、透明")

    add_heading2(doc, "6.2 合规认证")
    add_bullet(doc, "国内首批通过信通院 Claw 可信能力评估（五个检验维度，57 个具体能力项）")
    add_bullet(doc, "信通院智能原生编码 IDE 评估 4+ 级（同源 CodeBuddy，4类检测项目、13个能力项、177个能力子项评测表现优异）")
    add_bullet(doc, "同源工具 CodeBuddy 自举验证（在腾讯内部广泛使用，CodeBuddy 新增代码 100% 由 AI 自举生成）")

    # ===== 7. CodeBuddy 同源架构 =====
    add_heading2(doc, "七、CodeBuddy 同源架构")
    add_body(doc, "WorkBuddy 基于腾讯云代码助手 CodeBuddy 同源架构，CodeBuddy 有三种形态覆盖不同用户：")
    make_table(doc,
        ["产品形态", "适用用户", "核心特点"],
        [
            ["CodeBuddy IDE", "产品/设计师/全栈/初学者", '产设研一体工作台，\u201c对话即编程\u201d'],
            ["CodeBuddy 插件", "日常编码开发者", "即插即用、融入现有工作流"],
            ["CodeBuddy Code", "DevOps/运维/SRE/资深开发", "命令行工具，任务编排能力强"],
        ],
        [3, 4, 5]
    )

    # ===== 8. 素材索引 =====
    doc.add_page_break()
    add_heading1(doc, "八、素材索引")

    materials = [
        ("1", "WorkBuddy办公场景01_高效信息沟通_最佳实践案例", "https://csig.lexiangla.com/pages/e769ffa538b64e888380bf1659f58e9e"),
        ("2", "传统AI聊天_vs_WorkBuddy干活_最佳实践案例", "https://csig.lexiangla.com/pages/b96442a34bef4915b0e42ac078116d29"),
        ("3", "WorkBuddy-PM在日常工作应用场景实践-v5", "https://csig.lexiangla.com/pages/41aa17715ff64c15a8bd662ab8963cff"),
        ("4", "TAPD + 乐享 + WorkBuddy 联动方案", "https://csig.lexiangla.com/pages/98928c2886ac42e58d289bd6c8878a8d"),
        ("5", "WorkBuddy企业版主线逻辑说明材料", "https://csig.lexiangla.com/pages/dd473bab87ae46ab91a1049a101d35b4"),
        ("6", "WorkBuddy + 乐享AI知识库联合方案", "https://csig.lexiangla.com/pages/0be95da3d040443bacab438f7c678c52"),
        ("7", "workbuddy营销部门或市场部门场景", "https://csig.lexiangla.com/pages/ae018f44e3d64f96bfc6425df69463ed"),
        ("8", "腾讯企业版龙虾方案医疗行业", "https://csig.lexiangla.com/pages/ab39dd2b908944b9b29b3e30a7525c66"),
        ("9", "WorkBuddy 提效案例：构建个人学术信息助手", "https://csig.lexiangla.com/pages/e2cf5cbc5a8f44d28176829c92dbbe6e"),
        ("10", "【案例】新加坡-AI简历筛选提效90%", "https://csig.lexiangla.com/pages/bfce4bba7e754e38826b50396ced359d"),
        ("11", "基于乐享、企微、workbuddy等生成五粮液全民AI办公提效动员方案", "https://csig.lexiangla.com/pages/2b1e103954414fedad7c66094930b402"),
        ("12", "五粮液：AI助力质检提效解决方案", "https://csig.lexiangla.com/pages/13fe256cb6184810a790fc1e045247b1"),
        ("15", "WorkBuddy交付团队场景举例", "https://csig.lexiangla.com/pages/a704a6b92f864d40a92eefc363fd3dc6"),
    ]

    make_table(doc,
        ["序号", "素材名称", "链接"],
        [[m[0], m[1], m[2]] for m in materials],
        [1.5, 7, 5]
    )

    # Save
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "WorkBuddy产品介绍_场景与案例.docx")
    doc.save(output_path)
    print(f"DOCX saved to: {output_path}")

if __name__ == "__main__":
    main()
